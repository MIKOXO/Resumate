import io
import unittest

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt

import formatter


class FormatterTests(unittest.TestCase):
    def _document_bytes(self, document):
        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    def test_detects_body_size_inherited_from_normal_style(self):
        document = Document()
        document.styles["Normal"].font.name = "Arial"
        document.styles["Normal"].font.size = Pt(10.5)
        document.add_paragraph("First body paragraph with inherited formatting.")
        document.add_paragraph("Second body paragraph with inherited formatting.")

        font_name, font_size = formatter.detect_body_style(document)

        self.assertEqual(font_name, "Arial")
        self.assertEqual(font_size, Pt(10.5))

    def test_copies_dominant_header_and_body_paragraph_formatting(self):
        document = Document()
        document.styles["Normal"].font.name = "Aptos"
        document.styles["Normal"].font.size = Pt(10)
        document.add_paragraph("A long body paragraph establishes the normal resume style.")
        document.add_paragraph("Another body paragraph makes that style dominant.")
        header = document.add_paragraph("EXPERIENCE")
        header.paragraph_format.space_before = Pt(12)
        header.paragraph_format.space_after = Pt(4)
        header.paragraph_format.left_indent = Pt(18)
        header_run = header.runs[0]
        header_run.font.name = "Aptos Display"
        header_run.font.size = Pt(14)
        header_run.font.bold = True

        result = Document(io.BytesIO(formatter.insert_core_competencies(
            self._document_bytes(document), ["Relevant skill"], "resume.docx"
        )))
        core_header, bullet = result.paragraphs[-2:]

        self.assertEqual(core_header.text, "CORE COMPETENCIES")
        self.assertEqual(core_header.runs[0].font.name, "Aptos Display")
        self.assertEqual(core_header.runs[0].font.size, Pt(14))
        self.assertTrue(core_header.runs[0].font.bold)
        self.assertEqual(core_header.paragraph_format.space_before, Pt(12))
        self.assertEqual(core_header.paragraph_format.left_indent, Pt(18))
        self.assertTrue(core_header.paragraph_format.page_break_before)
        self.assertEqual(bullet.text, "• Relevant skill")
        self.assertEqual(formatter._paragraph_value(bullet, result, "size"), Pt(10))

    def test_title_case_headers_keep_title_case_section_name(self):
        document = Document()
        document.add_paragraph("Body text that establishes the default style.")
        header = document.add_paragraph("Professional Summary")
        header.runs[0].font.bold = True
        header.runs[0].font.size = Pt(13)

        result = Document(io.BytesIO(formatter.insert_core_competencies(
            self._document_bytes(document), ["Relevant skill"], "resume.docx"
        )))

        self.assertEqual(result.paragraphs[-2].text, "Core Competencies")

    def test_line_spacing_copied_to_bullets(self):
        document = Document()
        document.styles["Normal"].font.name = "Calibri"
        document.styles["Normal"].font.size = Pt(11)
        body = document.add_paragraph("Body paragraph with specific line spacing.")
        body.paragraph_format.line_spacing = 1.5
        body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        document.add_paragraph("Another body paragraph.")

        result = Document(io.BytesIO(formatter.insert_core_competencies(
            self._document_bytes(document), ["Skill one", "Skill two"], "resume.docx"
        )))
        bullet = result.paragraphs[-1]

        self.assertEqual(bullet.paragraph_format.line_spacing, 1.5)
        self.assertEqual(bullet.paragraph_format.line_spacing_rule, WD_LINE_SPACING.ONE_POINT_FIVE)

    def test_space_after_copied_to_bullets(self):
        document = Document()
        document.styles["Normal"].font.name = "Calibri"
        document.styles["Normal"].font.size = Pt(10)
        body = document.add_paragraph("Body paragraph with space after.")
        body.paragraph_format.space_after = Pt(6)
        document.add_paragraph("Another body paragraph.")

        result = Document(io.BytesIO(formatter.insert_core_competencies(
            self._document_bytes(document), ["Skill one"], "resume.docx"
        )))
        bullet = result.paragraphs[-1]

        self.assertEqual(bullet.paragraph_format.space_after, Pt(6))

    def test_local_spacing_preferred_over_global(self):
        document = Document()
        document.styles["Normal"].font.name = "Calibri"
        document.styles["Normal"].font.size = Pt(10)
        early = document.add_paragraph("Early paragraph with tight spacing.")
        early.paragraph_format.space_after = Pt(2)
        early.paragraph_format.line_spacing = 1.0
        early.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        for _ in range(6):
            p = document.add_paragraph("Later body paragraph with different spacing.")
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

        result = Document(io.BytesIO(formatter.insert_core_competencies(
            self._document_bytes(document), ["Skill one"], "resume.docx"
        )))
        bullet = result.paragraphs[-1]

        self.assertEqual(bullet.paragraph_format.space_after, Pt(8))
        self.assertEqual(bullet.paragraph_format.line_spacing, 1.15)

    def test_inherited_spacing_resolved_from_style(self):
        document = Document()
        document.styles["Normal"].font.name = "Calibri"
        document.styles["Normal"].font.size = Pt(10)
        document.styles["Normal"].paragraph_format.line_spacing = 1.15
        document.styles["Normal"].paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        document.styles["Normal"].paragraph_format.space_after = Pt(4)
        document.add_paragraph("Body paragraph inheriting from Normal style.")
        document.add_paragraph("Another paragraph inheriting spacing.")

        result = Document(io.BytesIO(formatter.insert_core_competencies(
            self._document_bytes(document), ["Skill one"], "resume.docx"
        )))
        bullet = result.paragraphs[-1]

        self.assertEqual(bullet.paragraph_format.line_spacing, 1.15)
        self.assertEqual(bullet.paragraph_format.space_after, Pt(4))

    def test_justified_alignment_matched(self):
        document = Document()
        document.styles["Normal"].font.name = "Calibri"
        document.styles["Normal"].font.size = Pt(10)
        for _ in range(3):
            p = document.add_paragraph("Justified body paragraph with enough text to fill.")
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        result = Document(io.BytesIO(formatter.insert_core_competencies(
            self._document_bytes(document), ["Skill one"], "resume.docx"
        )))
        bullet = result.paragraphs[-1]

        self.assertEqual(bullet.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)

    def test_left_indent_matched(self):
        document = Document()
        document.styles["Normal"].font.name = "Calibri"
        document.styles["Normal"].font.size = Pt(10)
        for _ in range(3):
            p = document.add_paragraph("Indented body paragraph.")
            p.paragraph_format.left_indent = Pt(24)

        result = Document(io.BytesIO(formatter.insert_core_competencies(
            self._document_bytes(document), ["Skill one"], "resume.docx"
        )))
        bullet = result.paragraphs[-1]

        self.assertEqual(bullet.paragraph_format.left_indent, Pt(24))

    def test_numbering_properties_stripped_from_bullets(self):
        document = Document()
        document.styles["Normal"].font.name = "Calibri"
        document.styles["Normal"].font.size = Pt(10)
        document.add_paragraph("Body paragraph one.")
        document.add_paragraph("Body paragraph two.")

        result = Document(io.BytesIO(formatter.insert_core_competencies(
            self._document_bytes(document), ["Skill one"], "resume.docx"
        )))
        bullet = result.paragraphs[-1]

        numPr = bullet._p.pPr.find(formatter.qn("w:numPr")) if bullet._p.pPr is not None else None
        self.assertIsNone(numPr)


if __name__ == "__main__":
    unittest.main()
