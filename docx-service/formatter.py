import io
import logging
from collections import Counter
from copy import deepcopy

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

KNOWN_HEADER_LABELS = {
    "experience", "professional experience", "work experience", "employment history",
    "education", "skills", "technical skills", "key skills", "summary",
    "professional summary", "qualifications", "certifications", "projects",
}
_MAX_HEADER_LEN = 50


def _paragraphs(doc):
    """Yield body paragraphs, including paragraphs nested in resume tables."""
    for paragraph in doc.paragraphs:
        yield paragraph

    def table_paragraphs(table):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
                for nested_table in cell.tables:
                    yield from table_paragraphs(nested_table)

    for table in doc.tables:
        yield from table_paragraphs(table)


def _style_value(style, property_name):
    while style:
        value = getattr(style.font, property_name)
        if value is not None:
            return value
        style = style.base_style
    return None


def _document_default_value(doc, property_name):
    """Read Word's document default when no run or named style defines it."""
    value = _style_value(doc.styles["Normal"], property_name)
    if value is not None:
        return value

    tag_by_property = {"size": "w:sz", "name": "w:rFonts"}
    tag = tag_by_property.get(property_name)
    if not tag:
        return None
    elements = doc.styles.element.xpath(f".//w:docDefaults/w:rPrDefault/w:rPr/{tag}")
    if not elements:
        return None
    if property_name == "size":
        value = elements[0].get(qn("w:val"))
        return Pt(int(value) / 2) if value and value.isdigit() else None
    return elements[0].get(qn("w:ascii")) or elements[0].get(qn("w:hAnsi"))


def _effective_run_value(run, paragraph, doc, property_name):
    value = getattr(run.font, property_name)
    if value is not None:
        return value
    value = _style_value(run.style, property_name)
    if value is not None:
        return value
    value = _style_value(paragraph.style, property_name)
    if value is not None:
        return value
    return _document_default_value(doc, property_name)


def _resolve_paragraph_attr(paragraph, attr):
    """Read a paragraph format attribute, resolving through the style hierarchy."""
    pf = paragraph.paragraph_format
    value = getattr(pf, attr)
    if value is not None:
        return value
    pstyle = paragraph.style
    if pstyle is not None:
        value = getattr(pstyle.paragraph_format, attr)
        if value is not None:
            return value
        base = pstyle.base_style
        if base is not None:
            value = getattr(base.paragraph_format, attr)
            if value is not None:
                return value
    return None


def _paragraph_spacing(paragraph):
    """Return the resolved spacing tuple for a paragraph."""
    return (
        _resolve_paragraph_attr(paragraph, "line_spacing"),
        _resolve_paragraph_attr(paragraph, "line_spacing_rule"),
        _resolve_paragraph_attr(paragraph, "space_before"),
        _resolve_paragraph_attr(paragraph, "space_after"),
    )


def _paragraph_alignment(paragraph):
    """Return the resolved alignment for a paragraph."""
    return _resolve_paragraph_attr(paragraph, "alignment")


def _paragraph_indentation(paragraph):
    """Return the resolved (left_indent, right_indent, first_line_indent) for a paragraph."""
    return (
        _resolve_paragraph_attr(paragraph, "left_indent"),
        _resolve_paragraph_attr(paragraph, "right_indent"),
        _resolve_paragraph_attr(paragraph, "first_line_indent"),
    )


def _meaningful_runs(paragraph):
    return [run for run in paragraph.runs if run.text.strip()]


def _paragraph_value(paragraph, doc, property_name):
    values = [
        _effective_run_value(run, paragraph, doc, property_name)
        for run in _meaningful_runs(paragraph)
    ]
    values = [value for value in values if value is not None]
    if values:
        return Counter(values).most_common(1)[0][0]
    return _style_value(paragraph.style, property_name) or _document_default_value(doc, property_name)


def _paragraph_is_bold(paragraph, doc):
    runs = _meaningful_runs(paragraph)
    return bool(runs) and all(
        _effective_run_value(run, paragraph, doc, "bold") is True for run in runs
    )


def _paragraph_all_caps(paragraph, doc):
    text = paragraph.text.strip()
    if text and text == text.upper() and any(character.isalpha() for character in text):
        return True
    return any(
        _effective_run_value(run, paragraph, doc, "all_caps") is True
        for run in _meaningful_runs(paragraph)
    )


def _normalized_label(text):
    return " ".join("".join(character if character.isalnum() else " " for character in text.lower()).split())


def _header_score(paragraph, doc, body_size):
    text = paragraph.text.strip()
    if not text or len(text) > _MAX_HEADER_LEN:
        return 0

    known_label = _normalized_label(text) in KNOWN_HEADER_LABELS
    heading_style = "heading" in (paragraph.style.name or "").lower()
    bold = _paragraph_is_bold(paragraph, doc)
    all_caps = _paragraph_all_caps(paragraph, doc)
    size = _paragraph_value(paragraph, doc, "size")
    larger_than_body = bool(size and body_size and size > body_size * 1.08)
    visual_signals = sum((bold, all_caps, larger_than_body))
    if not (known_label or heading_style) and visual_signals < 2:
        return 0
    return (4 if known_label else 0) + (2 if heading_style else 0) + visual_signals


def _style_signature(paragraph, doc):
    color = None
    for run in _meaningful_runs(paragraph):
        if run.font.color and run.font.color.type is not None:
            try:
                color = str(run.font.color.rgb)
            except Exception:
                pass
            break
    return (
        _paragraph_value(paragraph, doc, "name") or "Calibri",
        _paragraph_value(paragraph, doc, "size") or Pt(12),
        _paragraph_is_bold(paragraph, doc),
        color,
        _paragraph_all_caps(paragraph, doc),
        *_paragraph_spacing(paragraph),
    )


def _dominant_body_style(doc, excluded_paragraphs=(), paragraphs=None):
    excluded_ids = {id(paragraph) for paragraph in excluded_paragraphs}
    tally = Counter()
    representative = {}
    paragraph_iterable = paragraphs if paragraphs is not None else _paragraphs(doc)
    for paragraph in paragraph_iterable:
        text = paragraph.text.strip()
        if not text or id(paragraph) in excluded_ids:
            continue
        signature = _style_signature(paragraph, doc)[:2]
        tally[signature] += len(text)
        representative.setdefault(signature, paragraph)
    if not tally:
        return None, None, None
    signature, _ = tally.most_common(1)[0]
    return signature[0], signature[1], representative[signature]


def detect_body_style(doc, filename="<unknown>", paragraphs=None):
    """Return the dominant effective body font name and size."""
    font_name, font_size, _ = _dominant_body_style(doc, paragraphs=paragraphs)
    if font_name and font_size:
        return font_name, font_size
    logger.warning("Body fallback used for '%s': no body paragraphs found.", filename)
    return _document_default_value(doc, "name") or "Calibri", Pt(12)


def detect_header_style(doc, body_font_name, body_font_size, filename="<unknown>", paragraphs=None):
    """Return the dominant full header style and its representative paragraph."""
    paragraph_iterable = paragraphs if paragraphs is not None else _paragraphs(doc)
    candidates = [
        (paragraph, _header_score(paragraph, doc, body_font_size))
        for paragraph in paragraph_iterable
    ]
    candidates = [(paragraph, score) for paragraph, score in candidates if score]
    if not candidates:
        logger.warning("Header fallback used for '%s': no header candidates found.", filename)
        return {
            "font_name": body_font_name or "Calibri", "font_size": Pt(12),
            "bold": True, "color": None, "all_caps": False, "paragraph": None,
        }

    tally = Counter()
    representative = {}
    for paragraph, score in candidates:
        signature = _style_signature(paragraph, doc)
        tally[signature] += score
        representative.setdefault(signature, paragraph)
    signature, _ = tally.most_common(1)[0]
    return {
        "font_name": signature[0], "font_size": signature[1], "bold": signature[2],
        "color": signature[3], "all_caps": signature[4],
        "paragraph": representative[signature],
    }


def _copy_paragraph_format(source, target):
    source_properties = source._p.pPr
    if source_properties is None:
        return
    target_properties = target._p.pPr
    if target_properties is not None:
        target._p.remove(target_properties)
    target._p.insert(0, deepcopy(source_properties))


def _copy_run_format(source, target):
    source_properties = source._r.rPr
    if source_properties is None:
        return
    target_properties = target._r.rPr
    if target_properties is not None:
        target._r.remove(target_properties)
    target._r.insert(0, deepcopy(source_properties))


def _first_meaningful_run(paragraph):
    runs = _meaningful_runs(paragraph)
    return runs[0] if runs else None


def _apply_fallback_run_style(run, font_name, font_size, bold):
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold


def _local_body_spacing(doc, excluded_paragraphs=(), count=5, paragraphs=None):
    """Return the most common spacing tuple from the last N body paragraphs."""
    excluded_ids = {id(paragraph) for paragraph in excluded_paragraphs}
    paragraph_iterable = paragraphs if paragraphs is not None else _paragraphs(doc)
    body_paras = [
        paragraph for paragraph in paragraph_iterable
        if paragraph.text.strip() and id(paragraph) not in excluded_ids
    ]
    tail = body_paras[-count:] if len(body_paras) > count else body_paras
    if not tail:
        return None
    tally = Counter()
    for paragraph in tail:
        tally[_paragraph_spacing(paragraph)] += 1
    return tally.most_common(1)[0][0]


def _local_body_formatting(doc, excluded_paragraphs=(), count=5, paragraphs=None):
    """Return the most common (alignment, indentation) from the last N body paragraphs."""
    excluded_ids = {id(paragraph) for paragraph in excluded_paragraphs}
    paragraph_iterable = paragraphs if paragraphs is not None else _paragraphs(doc)
    body_paras = [
        paragraph for paragraph in paragraph_iterable
        if paragraph.text.strip() and id(paragraph) not in excluded_ids
    ]
    tail = body_paras[-count:] if len(body_paras) > count else body_paras
    if not tail:
        return None, None
    align_tally = Counter()
    indent_tally = Counter()
    for paragraph in tail:
        align_tally[_paragraph_alignment(paragraph)] += 1
        indent_tally[_paragraph_indentation(paragraph)] += 1
    return align_tally.most_common(1)[0][0], indent_tally.most_common(1)[0][0]


def _ensure_alignment(paragraph, alignment):
    """Set alignment on a paragraph if a value was detected."""
    if alignment is not None:
        paragraph.paragraph_format.alignment = alignment


def _ensure_indentation(paragraph, indentation):
    """Set left_indent, right_indent, first_line_indent on a paragraph."""
    if indentation is None:
        return
    left, right, first_line = indentation
    pf = paragraph.paragraph_format
    if left is not None:
        pf.left_indent = left
    if right is not None:
        pf.right_indent = right
    if first_line is not None:
        pf.first_line_indent = first_line


def _ensure_spacing(paragraph, spacing_tuple):
    """Set spacing on a paragraph from the detected profile."""
    line_spacing, line_rule, space_before, space_after = spacing_tuple
    pf = paragraph.paragraph_format
    if line_spacing is not None:
        pf.line_spacing = line_spacing
    if line_rule is not None:
        pf.line_spacing_rule = line_rule
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after


def _strip_numbering(pPr):
    """Remove w:numPr from a pPr element so copied list properties don't bleed through."""
    if pPr is None:
        return
    numPr = pPr.find(qn("w:numPr"))
    if numPr is not None:
        pPr.remove(numPr)


def insert_core_competencies(docx_bytes, bullet_lines, filename="<unknown>"):
    """Append a formatting-matched Core Competencies section to a .docx file."""
    doc = Document(io.BytesIO(docx_bytes))
    paragraphs = list(_paragraphs(doc))

    initial_body_name, initial_body_size = detect_body_style(doc, filename, paragraphs)
    header_style = detect_header_style(doc, initial_body_name, initial_body_size, filename, paragraphs)
    header_paragraph = header_style["paragraph"]
    _, _, body_paragraph = _dominant_body_style(
        doc,
        [header_paragraph] if header_paragraph else [],
        paragraphs,
    )
    body_font_name, body_font_size = initial_body_name, initial_body_size
    if body_paragraph:
        body_font_name, body_font_size = _style_signature(body_paragraph, doc)[:2]

    # Detect spacing, alignment, and indentation from the paragraphs closest to the insertion point.
    body_spacing = _local_body_spacing(doc, [header_paragraph] if header_paragraph else (), paragraphs=paragraphs)
    body_alignment, body_indentation = _local_body_formatting(
        doc,
        [header_paragraph] if header_paragraph else (),
        paragraphs=paragraphs,
    )
    header_spacing = _paragraph_spacing(header_paragraph) if header_paragraph else None
    header_alignment = _paragraph_alignment(header_paragraph) if header_paragraph else None
    header_indentation = _paragraph_indentation(header_paragraph) if header_paragraph else None

    header_para = doc.add_paragraph()
    if header_paragraph:
        _copy_paragraph_format(header_paragraph, header_para)
    if header_spacing:
        _ensure_spacing(header_para, header_spacing)
    if header_alignment is not None:
        _ensure_alignment(header_para, header_alignment)
    if header_indentation is not None:
        _ensure_indentation(header_para, header_indentation)
    # Keep the source header's horizontal layout while starting this section cleanly.
    header_para.paragraph_format.page_break_before = True
    header_run = header_para.add_run(
        "CORE COMPETENCIES" if header_style["all_caps"] else "Core Competencies"
    )
    header_source_run = _first_meaningful_run(header_paragraph) if header_paragraph else None
    if header_source_run:
        _copy_run_format(header_source_run, header_run)
    else:
        _apply_fallback_run_style(
            header_run, header_style["font_name"], header_style["font_size"], header_style["bold"]
        )

    body_source_run = _first_meaningful_run(body_paragraph) if body_paragraph else None
    for line in bullet_lines:
        line = line.strip()
        if not line:
            continue
        bullet_para = doc.add_paragraph()
        if body_paragraph:
            _copy_paragraph_format(body_paragraph, bullet_para)
            _strip_numbering(bullet_para._p.pPr)
        if body_spacing:
            _ensure_spacing(bullet_para, body_spacing)
        if body_alignment is not None:
            _ensure_alignment(bullet_para, body_alignment)
        if body_indentation is not None:
            _ensure_indentation(bullet_para, body_indentation)
        bullet_run = bullet_para.add_run(f"\u2022 {line}")
        if body_source_run:
            _copy_run_format(body_source_run, bullet_run)
        else:
            _apply_fallback_run_style(bullet_run, body_font_name, body_font_size, False)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
